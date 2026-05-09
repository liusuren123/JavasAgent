"""Word 文档操作子模块。

提供 Word (.docx) 文档的创建、读取、追加、编辑能力。
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from loguru import logger

from src.utils.path_safety import PathSafetyError, safe_resolve_path


# ======================================================================
# 公共辅助
# ======================================================================


def _safe_path(workspace: Path, user_path: str, *, allow_create_parents: bool = False) -> Path:
    """安全解析用户路径。"""
    return safe_resolve_path(workspace, user_path, allow_create_parents=allow_create_parents)


def _check_docx_import() -> str | None:
    """检查 python-docx 是否已安装，未安装则返回错误信息。"""
    try:
        import docx  # noqa: F401
        return None
    except ImportError:
        return "python-docx 未安装，请运行: pip install python-docx"


# ======================================================================
# 操作实现
# ======================================================================


async def read_docx(workspace: Path, params: dict) -> dict:
    """读取 Word 文档内容。

    Params:
        path: 文件路径（相对于 workspace）
        max_paragraphs: 最大读取段落数（默认 500）
    """
    try:
        path = _safe_path(workspace, params["path"])
    except PathSafetyError as e:
        return {"error": str(e)}

    if not path.exists():
        return {"error": f"文件不存在: {path}"}

    err = _check_docx_import()
    if err:
        return {"error": err}

    try:
        from docx import Document

        doc = Document(str(path))
        max_para = params.get("max_paragraphs", 500)

        paragraphs: list[dict[str, Any]] = []
        for i, para in enumerate(doc.paragraphs[:max_para]):
            if para.text.strip():
                paragraphs.append({
                    "index": i,
                    "style": para.style.name if para.style else "Normal",
                    "text": para.text,
                })

        # 提取表格内容
        tables: list[list[list[str]]] = []
        for table in doc.tables:
            rows_data: list[list[str]] = []
            for row in table.rows:
                rows_data.append([cell.text for cell in row.cells])
            tables.append(rows_data)

        logger.info(f"读取 Word 文档: {path} ({len(paragraphs)} 段落, {len(tables)} 表格)")
        return {
            "path": str(path),
            "paragraphs": paragraphs,
            "paragraph_count": len(paragraphs),
            "tables": tables,
            "table_count": len(tables),
        }
    except Exception as e:
        logger.error(f"读取 Word 文档失败: {e}")
        return {"error": f"读取失败: {e}"}


async def create_docx(workspace: Path, params: dict) -> dict:
    """创建 Word 文档。

    Params:
        path: 文件路径
        title: 文档标题（可选）
        paragraphs: 段落文本列表
        headings: 标题列表 [{"level": 1, "text": "标题"}]（可选）
    """
    try:
        path = _safe_path(workspace, params["path"], allow_create_parents=True)
    except PathSafetyError as e:
        return {"error": str(e)}

    err = _check_docx_import()
    if err:
        return {"error": err}

    try:
        from docx import Document

        doc = Document()

        # 添加标题
        title = params.get("title")
        if title:
            doc.add_heading(title, level=0)

        # 添加标题列表
        for heading in params.get("headings", []):
            level = min(heading.get("level", 1), 4)
            doc.add_heading(heading["text"], level=level)

        # 添加段落
        for text in params.get("paragraphs", []):
            doc.add_paragraph(text)

        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))

        logger.info(f"创建 Word 文档: {path}")
        return {"path": str(path), "created": True}
    except Exception as e:
        logger.error(f"创建 Word 文档失败: {e}")
        return {"error": f"创建失败: {e}"}


async def append_docx(workspace: Path, params: dict) -> dict:
    """追加内容到现有 Word 文档。

    Params:
        path: 文件路径
        paragraphs: 要追加的段落文本列表
        headings: 要追加的标题列表（可选）
    """
    try:
        path = _safe_path(workspace, params["path"])
    except PathSafetyError as e:
        return {"error": str(e)}

    if not path.exists():
        return {"error": f"文件不存在: {path}"}

    err = _check_docx_import()
    if err:
        return {"error": err}

    try:
        from docx import Document

        doc = Document(str(path))

        for heading in params.get("headings", []):
            level = min(heading.get("level", 1), 4)
            doc.add_heading(heading["text"], level=level)

        for text in params.get("paragraphs", []):
            doc.add_paragraph(text)

        doc.save(str(path))

        added_count = len(params.get("paragraphs", [])) + len(params.get("headings", []))
        logger.info(f"追加 Word 文档: {path} (+{added_count} 段)")
        return {"path": str(path), "appended": added_count}
    except Exception as e:
        logger.error(f"追加 Word 文档失败: {e}")
        return {"error": f"追加失败: {e}"}


async def edit_docx(workspace: Path, params: dict) -> dict:
    """编辑已有 Word 文档中的文本。

    支持两种模式：
    1. 按段落索引替换：指定 ``replacements`` 列表，每项含 index + text
    2. 文本搜索替换：指定 ``search_replace`` 列表，每项含 search + replace

    Params:
        path: 文件路径
        replacements: 按索引替换列表 [{"index": 0, "text": "新文本"}]
        search_replace: 搜索替换列表 [{"search": "旧文本", "replace": "新文本"}]
    """
    try:
        path = _safe_path(workspace, params["path"])
    except PathSafetyError as e:
        return {"error": str(e)}

    if not path.exists():
        return {"error": f"文件不存在: {path}"}

    err = _check_docx_import()
    if err:
        return {"error": err}

    try:
        from docx import Document

        doc = Document(str(path))
        replaced = 0

        # 模式 1：按段落索引替换
        for item in params.get("replacements", []):
            idx = item["index"]
            new_text = item["text"]
            if 0 <= idx < len(doc.paragraphs):
                for run in doc.paragraphs[idx].runs:
                    run.text = ""
                # 清空后设置第一个 run 的文本；若无 run 则无法修改
                if doc.paragraphs[idx].runs:
                    doc.paragraphs[idx].runs[0].text = new_text
                else:
                    # 没有 runs 的段落直接设置 text（清空格式）
                    doc.paragraphs[idx].text = new_text
                replaced += 1

        # 模式 2：搜索替换
        for item in params.get("search_replace", []):
            search_text = item["search"]
            replace_text = item["replace"]
            for para in doc.paragraphs:
                if search_text in para.text:
                    for run in para.runs:
                        if search_text in run.text:
                            run.text = run.text.replace(search_text, replace_text)
                            replaced += 1
                            break

        doc.save(str(path))

        logger.info(f"编辑 Word 文档: {path} ({replaced} 处替换)")
        return {"path": str(path), "replaced": replaced}
    except Exception as e:
        logger.error(f"编辑 Word 文档失败: {e}")
        return {"error": f"编辑失败: {e}"}
